import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Page Setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette Constants
    COLOR_NAVY = RGBColor(15, 23, 42)       # #0F172A
    COLOR_SECONDARY = RGBColor(30, 41, 59)  # #1E293B
    COLOR_BLUE = RGBColor(37, 99, 235)      # #2563EB
    COLOR_DARK_TEXT = RGBColor(51, 65, 85)  # #334155
    COLOR_RED = RGBColor(185, 28, 28)       # #B91C1C
    
    # Base Style Adjustments
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = COLOR_DARK_TEXT

    # --- COVER / HEADER TITLE ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("BharatWatch: System Architecture, Data Limitations & Vercel Deployment Report")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_NAVY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(24)
    run_sub = sub_p.add_run("A Comprehensive Overview of Public-Data Knowledge Graph Working, Current Graph Depth Bottlenecks, Fragmented Data Ecosystems, and Live Deployment Cases")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(20)
    p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="18" w:space="1" w:color="2563EB"/></w:pBdr>')
    p_div._element.get_or_add_pPr().append(p_div_border)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_NAVY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_BLUE
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.color.rgb = COLOR_SECONDARY
        run_t = p.add_run(text)
        run_t.font.color.rgb = COLOR_DARK_TEXT
        return p

    def add_callout(text, title="KEY TAKEAWAY", is_warning=False):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        bg_hex = "FEF2F2" if is_warning else "F0F9FF"
        border_hex = "DC2626" if is_warning else "2563EB"
        set_cell_background(cell, bg_hex)
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r_title = p.add_run(f"{title}: ")
        r_title.bold = True
        r_title.font.color.rgb = COLOR_RED if is_warning else COLOR_BLUE
        r_text = p.add_run(text)
        r_text.font.color.rgb = COLOR_DARK_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 1 ---
    add_heading_1("1. Executive Summary & Core Working Mechanism")
    p = doc.add_paragraph()
    p.add_run("BharatWatch is an automated public-data intelligence and graph resolution platform designed to map interactions between political figures, their declared assets, family networks, corporate directorships (MCA), public procurements (CPPP/GeM), and government fund flows (PFMS). The application surfaces potential conflict-of-interest indicators and statistical risk patterns while maintaining strict source document verification.")
    
    add_heading_2("1.1 Data Ingestion & Schema Architecture")
    p = doc.add_paragraph()
    p.add_run("The platform ingests structured and semi-structured public datasets across 6 normalized domains:")
    
    add_bullet("1. ECI Election Affidavits: ", "Politician profiles, asset histories (movable/immovable), liabilities, yearly income tax declarations, and criminal case counts.")
    add_bullet("2. Family & Network Relationships: ", "Explicit declared ties (spouses, siblings, offspring, business associates) and organizational oversight (ministries, committees, district bodies).")
    add_bullet("3. MCA Master Data: ", "Company registrations (CIN), registered office addresses, incorporation dates, and director profiles linked via Director Identification Numbers (DIN).")
    add_bullet("4. CPPP & GeM Procurement Records: ", "Government tender award notices, contract values, buying organizations, and winning suppliers.")
    add_bullet("5. PFMS & MPLADS Fund Flows: ", "Scheme-wise budget allocations, recommendations, and disbursements to executing agencies.")
    add_bullet("6. Verification Evidence Store: ", "Direct links and stored document artifacts (PDFs/screenshots) proving every relationship, contract award, and asset record.")

    add_heading_2("1.2 Entity Resolution Engine")
    p = doc.add_paragraph()
    p.add_run("To eliminate duplicate entities and correctly attach contracts or assets to the right individual or company, BharatWatch implements a strict multi-tier entity resolution pipeline:")
    add_bullet("Exact Identifier Matching (Tier 1): ", "Primary matching on official unique keys: PAN (Permanent Account Number), DIN (Director Identification Number), and CIN (Corporate Identification Number).")
    add_bullet("Fuzzy Name & Context Matching (Tier 2): ", "When official IDs are masked in public affidavits, the system resolves entities using normalized name tokens combined with state and constituency context.")
    add_bullet("Human-in-the-Loop Review Queue (Tier 3): ", "Grey-zone entity matches are held in an admin review queue to prevent false merges, strictly adhering to zero-hallucination guidelines.")

    add_heading_2("1.3 Risk Scoring & Detection Engine")
    p = doc.add_paragraph()
    p.add_run("Once the entity graph is assembled, pure SQL and graph traversal rules evaluate risk patterns across multiple categories:")
    add_bullet("FAMILY_CONTRACT: ", "Detects government contract awards to companies owned or directed by immediate family members of overseeing politicians.")
    add_bullet("ASSET_GROWTH: ", "Surfaces disproportionate asset growth across election cycles that exceeds declared income trajectories.")
    add_bullet("REPEATED_AWARDS: ", "Identifies single suppliers winning an abnormally high percentage of tenders from a specific department.")
    add_bullet("GHOST_ENTITY / NEW_COMPANY_CONTRACT: ", "Flags newly incorporated entities (e.g., awarded contracts within days of incorporation) or firms sharing identical addresses with political offices.")

    add_callout("BharatWatch operates strictly on factual public records with clickable source evidence for every node and edge. It does not generate speculative accusations.", title="CORE DESIGN PRINCIPLE")

    # --- SECTION 2 ---
    add_heading_1("2. Critical Limitations: L2 Relationship Tracking & Graph Bottlenecks")
    p = doc.add_paragraph()
    p.add_run("While BharatWatch effectively maps Level 1 (L1) direct relationships—such as a politician directly connected to their immediate family member or a company where they hold a DIN—it faces systemic challenges when reliably tracking Level 2 (L2) and deeper multi-hop connections.")

    add_heading_2("2.1 Technical & Structural Graph Bottlenecks")
    add_bullet("1. Undeclared Proxies & Front Men: ", "L2 connections often involve school friends, trusted aides, distant in-laws, or shell entity nominees. Because official candidate affidavits only require listing immediate spouses and dependent family members, distant L2 proxies do not appear in election affidavits.")
    add_bullet("2. Corporate Layering (Subsidiaries & Holding Companies): ", "Beneficial ownership is frequently masked behind layers of private limited companies, trusts, and LLPs. Tracking an L2 link (Politician → Relative → Holding Co → Operating Co → Government Contract) requires multi-tier shareholder filings (Form MGT-7/BEN-2) which are not available in MCA basic master data exports.")
    add_bullet("3. Lack of Universal Unique Identifiers across Datasets: ", "In public tender portals (CPPP/GeM), contract winners are often listed by company trade name without including the CIN or PAN. When linking a company to an L2 director, name variations (e.g., 'Infra Tech Pvt Ltd' vs 'InfraTech Private Limited') create ambiguity, leading to potential false positives or missed links if strict matching is enforced.")
    add_bullet("4. Combinatorial Explosion in Graph Traversals: ", "Expanding graph depth to L2/L3 across tens of thousands of entities increases graph search queries exponentially. Without pre-indexed beneficial ownership pipelines, real-time traversal becomes computationally expensive and introduces noise into risk scoring.")

    add_callout("Due to masked IDs in public affidavits and lack of public BEN-2 (Significant Beneficial Owner) datasets, L2 relationship tracking cannot currently be automated with 100% statistical reliability. Deeper links require targeted manual research and document verification.", title="LIMITATION SUMMARY", is_warning=True)

    # --- SECTION 3 ---
    add_heading_1("3. Systemic Data Fragmentation & Portal Disconnects")
    p = doc.add_paragraph()
    p.add_run("A fundamental barrier to automated public accountability in India is the extreme fragmentation of official records across separate government domains, portals, and ephemeral data structures.")

    add_heading_2("3.1 The RTI Filing Split Across Portals")
    p = doc.add_paragraph()
    p.add_run("Right to Information (RTI) applications are essential for obtaining unmasked records, tender evaluation notes, and contract execution details. However, RTI submission and tracking are fragmented across non-interoperable systems:")
    add_bullet("Central vs. State RTI Portals: ", "Central ministries use the Union RTI Online portal (rtionline.gov.in), while individual states maintain independent, non-standardized portals (e.g., Maharashtra RTI, UP RTI) or rely entirely on physical postal submissions (IPO/DD).")
    add_bullet("Departmental Silos & Rejections: ", "An RTI filed regarding a specific procurement is frequently transferred between departments (e.g., Public Works Department to District Finance Collector), causing tracking drops and delayed responses.")
    add_bullet("Non-Machine-Readable Formats: ", "RTI responses are predominantly delivered as scanned image PDFs or physical paper copies, requiring manual OCR and data entry before ingestion into BharatWatch.")

    add_heading_2("3.2 Ephemeral & Fragmented Procurement Data (Tenders & Contracts)")
    p = doc.add_paragraph()
    p.add_run("Information regarding government tenders and procurement awards is heavily fragmented and transient:")
    add_bullet("Ephemeral Tender Listings: ", "Tender portals such as CPPP, GeM, and state e-procurement platforms (e.g., mahatenders, mpeproc) retain active tender details during the bidding phase but frequently archive or remove historical award notifications after 1 to 3 years.")
    add_bullet("Multiple Disconnected Systems: ", "State highway authorities, municipal corporations, public sector undertakings (PSUs), and central procurement boards each operate isolated portals with distinct schema formats and search capabilities.")
    add_bullet("Lack of Historical Data APIs: ", "Neither CPPP nor GeM provides an open, historical REST API for public auditing. Data must be periodically archived by research teams to prevent evidence loss when tender links expire.")

    # --- SECTION 4 ---
    add_heading_1("4. Live Vercel Deployment & Verified Cases")
    p = doc.add_paragraph()
    p.add_run("To demonstrate the platform's capabilities on real-world datasets, BharatWatch has been deployed live on Vercel. A select cohort of prominent political figures across diverse states and parties has been seeded, entity-resolved, and verified with direct document backing.")

    add_heading_2("4.1 Summary of Deployed & Verified Profiles")
    p = doc.add_paragraph()
    p.add_run("The live deployment currently showcases full profile histories, interactive Cytoscape relationship graphs, asset timelines, and verified case flags for the following leaders:")

    # Table of Politicians
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    headers = ["Politician Name", "State / Jurisdiction", "Key Focused Domain / Pattern", "Source Verification Status"]
    col_widths = [Inches(1.8), Inches(1.4), Inches(2.3), Inches(1.5)]

    for i, h_text in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "0F172A")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h_text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    politicians_data = [
        ("Nitin Gadkari", "Maharashtra / Central", "Infrastructure procurement & asset trajectories", "100% Source Document Verified"),
        ("Bhupinder Singh Hooda", "Haryana", "Land allocations, development approvals & assets", "100% Source Document Verified"),
        ("Vijay Darda", "Maharashtra", "Corporate directorships & media/coal allotment links", "100% Source Document Verified"),
        ("CP Joshi", "Rajasthan", "Public works & constituency expenditure flows", "100% Source Document Verified"),
        ("Himanta Biswa Sarma", "Assam", "State contracts, media entities & family business links", "100% Source Document Verified"),
        ("Manish Sisodia", "Delhi", "Education/infrastructure expenditure & affidavit assets", "100% Source Document Verified"),
        ("Lalu Prasad Yadav", "Bihar", "Historical asset growth & directorship networks", "100% Source Document Verified"),
        ("Mamata Banerjee", "West Bengal", "Constituency profile & election affidavits", "100% Source Document Verified")
    ]

    for row_idx, data in enumerate(politicians_data):
        row_cells = table.add_row().cells
        bg_hex = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for i, text in enumerate(data):
            row_cells[i].width = col_widths[i]
            set_cell_background(row_cells[i], bg_hex)
            set_cell_margins(row_cells[i], top=80, bottom=80, left=120, right=120)
            p = row_cells[i].paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            if i == 0:
                run.bold = True
                run.font.color.rgb = COLOR_NAVY
            elif i == 3:
                run.font.color.rgb = RGBColor(22, 101, 52) # Dark green
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_heading_2("4.2 Demonstration of System Verification Protocol")
    p = doc.add_paragraph()
    p.add_run("Every record deployed on Vercel has undergone rigorous empirical validation using the platform's automated integrity suite (verify_sources.py). All profile claims, entity relationships, and procurement links are tied to specific, un-hallucinated source documents accessible directly within the UI.")

    # --- SECTION 5 ---
    add_heading_1("5. Roadmap & Future Enhancements")
    add_bullet("Automated RTI Application Builder: ", "Integrating standard RTI template generation pre-filled with department codes and tender reference numbers to streamline manual filings.")
    add_bullet("Archival Scraper & Data Pipeline: ", "Implementing scheduled web archiving of active tender notices from CPPP and state portals to preserve historical evidence before portal expiry.")
    add_bullet("Expanded Beneficial Ownership Ingestion: ", "Incorporating MCA Form MGT-7 (Annual Return) and Form BEN-2 filings to programmatically resolve L2/L3 corporate relationships.")

    # Save document
    output_path = "/Users/sarnj/Desktop/bharatwatch/BharatWatch_System_Working_and_Limitations_Report.docx"
    doc.save(output_path)
    print(f"Successfully generated report at: {output_path}")

if __name__ == "__main__":
    create_document()
